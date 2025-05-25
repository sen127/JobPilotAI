def sanitize(text):
    return (
        text.replace(" ", "")
            .replace("/", "_")
            .replace("\\", "_")
            .replace(":", "")
            .replace("?", "")
            .replace("|", "")
            .replace('"', "")
            .replace("'", "")
            .strip() or "Unknown"
    )

def copy_template_and_insert_letter(template_path, body_text, disclaimer):
    doc = Document(template_path)
    new_doc = Document()
    paras = doc.paragraphs
    greeting_idx = next((i for i, p in enumerate(paras) if p.text.strip().startswith("Dear")), 0)
    closing_idx = next((i for i, p in reversed(list(enumerate(paras))) if "Sincerely" in p.text or "," in p.text), len(paras)-1)

    new_doc.add_paragraph(disclaimer)
    for i in range(greeting_idx):
        new_doc.add_paragraph(paras[i].text, style=paras[i].style)
    new_doc.add_paragraph(paras[greeting_idx].text, style=paras[greeting_idx].style)
    for line in body_text.strip().split("\n"):
        if line.strip():
            new_doc.add_paragraph(line.strip())
    for i in range(closing_idx, len(paras)):
        new_doc.add_paragraph(paras[i].text, style=paras[i].style)
    return new_doc


import os
import json
import time
from pathlib import Path
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from backend.llm_writer import generate_cover_letter, rewrite_resume_sentences
from backend.tracker import log_status
from backend.user_session import init_user_session
from docx import Document
from datetime import datetime

def apply_to_job(job_file, user):
    print(f"\n🚀 Applying to: {job_file.name}")
    with open(job_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    job_data = {line.split(":")[0].strip(): line.split(":", 1)[1].strip() for line in lines if ":" in line}
    job_url = job_data.get("Job URL", "")

    job_title = job_data.get("Job Title", "Unknown Role")
    company = job_data.get("Company", "Unknown Company")
    job_description_raw = json.dumps(job_data, indent=2)
    job_description = job_description_raw[:3000] + ("\n...(truncated)" if len(job_description_raw) > 3000 else "")

    if "(To be fetched manually" in job_description:
        print("⚠️ Warning: Job description is a placeholder.")

    for required in ["prompt_path", "resume_path", "cover_template_path"]:
        if not Path(user[required]).exists():
            print(f"❌ Missing required file: {user[required]}")
            return

    # Generate updated resume
    resume_doc = Document(user["resume_path"])
    resume_lines = [p.text.strip() for p in resume_doc.paragraphs if p.text.strip()]
    updated_resume_lines = rewrite_resume_sentences(resume_lines, job_description)
    updated_resume_doc = Document()
    for orig_para, new_text in zip(resume_doc.paragraphs, updated_resume_lines):
        if new_text.strip():
            para = updated_resume_doc.add_paragraph(new_text)
            para.style = orig_para.style

    # Save resume
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    resume_filename = f"resume_{sanitize(company)}_{sanitize(job_title)}_{user['username']}_{timestamp}.docx"
    resume_path = Path(user["user_dir"]) / resume_filename
    updated_resume_doc.save(resume_path)

    user_background = "\n".join(resume_lines)

    # Generate cover letter
    try:
        with open(user["prompt_path"], "r") as f:
            system_prompt = f.read().strip()
        if not system_prompt:
            raise ValueError("Prompt file is empty.")
    except Exception as e:
        print(f"❌ Failed to load prompt: {e}")
        log_status(user["username"], job_file.name, "FAILED", f"Missing or invalid prompt: {e}")
        return

    cover_letter = generate_cover_letter(job_title, company, job_description, user_background, system_prompt)

    if not cover_letter:
        print("❌ Failed to generate cover letter.")
        log_status(user["username"], job_file.name, "FAILED", "Cover letter generation failed")
        return

    cover_doc = copy_template_and_insert_letter(
        user["cover_template_path"],
        cover_letter,
        "This job was applied using an AI bot I built to automate job applications and demonstrate my product thinking, engineering skills, and initiative."
    )

    # Save cover letter
    cover_filename = f"cover_letter_{sanitize(company)}_{sanitize(job_title)}_{user['username']}_{timestamp}.docx"
    cover_path = Path(user["user_dir"]) / cover_filename
    cover_doc.save(cover_path)

    print(f"📄 Resume: {resume_path.name}")
    print(f"📄 Cover Letter: {cover_path.name}")

    if not job_url or job_url == "N/A":
        print("❌ No valid job URL found.")
        log_status(user["username"], job_file.name, "FAILED", "Missing job URL")
        return

    # Launch browser
    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    chrome_options.add_argument("--user-data-dir=/tmp/chrome-bhim-auto")
    chrome_options.add_argument("--profile-directory=Default")
    chrome_options.add_experimental_option("detach", False)

    try:
        driver = webdriver.Chrome(service=Service(), options=chrome_options)
        driver.get(job_url)

        # Wait for page to load or cookie accept
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(3)

        # Placeholder: logic to click "Apply" and auto-fill form goes here
        print("🛠️ Simulating job application...")

        log_status(user["username"], job_file.name, "APPLIED", f"{job_url} | Resume: {resume_path.name}, Cover: {cover_path.name}")
        driver.quit()

    except Exception as e:
        print(f"❌ Error during application: {e}")
        log_status(user["username"], job_file.name, "FAILED", str(e))

def main():
    from backend.tracker import already_applied
    user = init_user_session(include_paths=True)
    scraped_dir = Path(f"data/users/{user['username']}/scraped_jobs")
    job_files = list(scraped_dir.glob("*.txt"))

    if not job_files:
        print("⚠️ No job files found.")
        return

    for job_file in job_files:
        if already_applied(user["username"], job_file.name):
            print(f"⏭️ Already applied to: {job_file.name}")
            continue
        apply_to_job(job_file, user)

if __name__ == "__main__":
    main()