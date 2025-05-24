from docx import Document
from backend.llm_writer import generate_cover_letter
from datetime import datetime
from pathlib import Path

def replace_middle_paragraphs_with_preserved_formatting(template_path, new_body_text, disclaimer):
    doc = Document(template_path)
    new_doc = Document()

    # Add header disclaimer
    new_doc.add_paragraph(disclaimer)

    paras = doc.paragraphs
    greeting_idx = next((i for i, p in enumerate(paras) if p.text.strip().startswith("Dear")), 0)
    closing_idx = next((i for i, p in reversed(list(enumerate(paras))) if "Sincerely" in p.text or "," in p.text), len(paras)-1)

    for i in range(greeting_idx):
        new_doc.add_paragraph(paras[i].text, style=paras[i].style)

    new_doc.add_paragraph(paras[greeting_idx].text, style=paras[greeting_idx].style)

    for line in new_body_text.strip().split("\n"):
        if line.strip():
            new_doc.add_paragraph(line.strip())

    for i in range(closing_idx, len(paras)):
        new_doc.add_paragraph(paras[i].text, style=paras[i].style)

    return new_doc


def generate_cover_letter_docx(job_title, company_name, job_description, user_background, system_prompt, template_path, save_dir, username):
    disclaimer = "This job was applied using an AI bot I built to automate job applications and demonstrate my product thinking, engineering skills, and initiative."
    letter = generate_cover_letter(job_title, company_name, job_description, user_background, system_prompt)
    if not letter:
        print("❌ GPT failed to generate cover letter.")
        return ""

    final_doc = replace_middle_paragraphs_with_preserved_formatting(template_path, letter, disclaimer)

    filename = f"cover_letter_{company_name.replace(' ', '')}_{job_title.replace(' ', '')}_{username}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    output_path = Path(save_dir) / filename
    final_doc.save(output_path)

    return output_path