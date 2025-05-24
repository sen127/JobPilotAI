from docx import Document
from backend.llm_writer import rewrite_resume_sentences
from pathlib import Path
from datetime import datetime

def generate_resume_docx(resume_path, job_description, username, company_name, job_title, save_dir):
    doc = Document(resume_path)

    lines = [para.text.strip() for para in doc.paragraphs]
    rewritten_lines = rewrite_resume_sentences(lines, job_description)

    for para, new_line in zip(doc.paragraphs, rewritten_lines):
        para.text = new_line

    filename = f"resume_{company_name.replace(' ', '')}_{job_title.replace(' ', '')}_{username}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    output_path = Path(save_dir) / filename
    doc.save(output_path)

    return output_path